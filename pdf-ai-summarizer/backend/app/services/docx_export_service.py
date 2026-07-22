import io

from docx import Document
from fastapi import HTTPException
from pathlib import Path
from pdf2docx import Converter

from app.core.config import settings
from app.services.pdf_reader import get_pdf_text
from app.services.path_service import _resolve_pdf_path


def convert_pdf_to_docx_mechanical(document_id: str) -> bytes:
    pdf_path = _resolve_pdf_path(document_id)
    buffer = io.BytesIO()
    try:
        converter = Converter(pdf_file=str(pdf_path))
        converter.convert(buffer)
        converter.close()
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Cannot convert PDF to DOCX") from exc

    return buffer.getvalue()

def convert_pdf_to_docx_ai(document_id: str, extraction_mode: str = "text") -> bytes:
    text = get_pdf_text(document_id, extraction_mode)
    document = Document()

    table_buffer: list[str] = []

    def flush_table_buffer() -> None:
        if len(table_buffer) < 2:
            table_buffer.clear()
            return

        header_cells = [cell.strip() for cell in table_buffer[0].strip("|").split("|")]
        data_rows = [
            [cell.strip() for cell in row.strip("|").split("|")]
            for row in table_buffer[2:]
        ]

        table = document.add_table(rows=1, cols=len(header_cells))
        table.style = "Light Grid Accent 1"
        for i, cell_text in enumerate(header_cells):
            table.rows[0].cells[i].text = cell_text

        for row in data_rows:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row):
                if i < len(row_cells):
                    row_cells[i].text = cell_text

        table_buffer.clear()

    for line in text.split("\n"):
        if line.strip().startswith("|"):
            table_buffer.append(line)
            continue

        if table_buffer:
            flush_table_buffer()

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.strip():
            document.add_paragraph(line)

    if table_buffer:
        flush_table_buffer()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
