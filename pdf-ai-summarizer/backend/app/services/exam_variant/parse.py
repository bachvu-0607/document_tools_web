import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.services.exam_variant.models import Option, Part, Question

# 1 o bang CHI co dung 1 chu cai, khong co gi khac (vd cau 1 trong bang: nhan
# "A" nam rieng 1 o, noi dung "It is modern and clean." nam o o ke ben).
BARE_LABEL_RE = re.compile(r"^[A-D]$")

# "I. LISTENING", "III. Mark the letter ...": so La Ma + dau cham + chu hoa.
# CHI liet ke I..X tuong minh (khong dung [IVXLCDM] chung chung) vi "C."/"D."
# cung la so La Ma hop le (100/500) nhung o day trung voi nhan dap an C/D.
PART_HEADER_RE = re.compile(r"^(?P<num>VIII|VII|VI|IV|IX|III|II|I|V|X)\.\s+[A-ZĐ]")
# "Question 1.", "Cau 12:"
QUESTION_RE = re.compile(r"^(?:Question|C[aâ]u)\s*(?P<num>\d+)[.:]?\s*", re.IGNORECASE)
# 1 dap an rieng 1 dong: "A. text" / "A   text" / "A) text"
# KHONG dung ":" o day - "A:"/"B:" trong bai doc hoi thoai la nhan NGUOI NOI,
# khong phai dap an trac nghiem (vd cau 23: "A: I'm going to throw...").
OPTION_RE = re.compile(r"^(?P<label>[A-D])[.)]?\s+(?P<rest>.+)$")
# nhieu dap an chung 1 dong, cach nhau >=2 khoang trang: "A. start   B. starts   C. will start"
INLINE_OPTION_RE = re.compile(r"(?P<label>[A-D])[.)]\s*(?P<rest>.*?)(?=\s{2,}[A-D][.)]|$)")

RED_R_MIN = 120
RED_DOMINANCE = 40
CHOICE_LETTERS = "ABCD"


def _has_auto_numbering(paragraph: Paragraph) -> bool:
    """True neu dong dau tien cua paragraph dung tinh nang danh so tu dong
    cua Word (numPr) - luc do chu "A." nguoi dung thay tren Word la do Word
    TU VE RA, khong nam trong noi dung text that (paragraph.text khong co)."""
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    return p_pr.find(qn("w:numPr")) is not None


def _runs_in_range(paragraph: Paragraph, start: int, end: int) -> list:
    """Tim cac run co text nam trong khoang [start, end) cua paragraph.text.
    Can thiet vi python-docx khong cho biet truc tiep run nao ung voi 1 doan
    text con (vi du 1 dap an trong dong nhieu dap an gop chung 1 paragraph).
    """
    result = []
    offset = 0
    for run in paragraph.runs:
        run_start, run_end = offset, offset + len(run.text)
        if run_start < end and run_end > start:
            result.append(run)
        offset = run_end
    return result


def _is_red_range(paragraph: Paragraph, start: int, end: int) -> bool:
    for run in _runs_in_range(paragraph, start, end):
        color = run.font.color
        if color is None or color.type is None or color.rgb is None:
            continue
        r, g, b = color.rgb[0], color.rgb[1], color.rgb[2]
        if r > RED_R_MIN and r > g + RED_DOMINANCE and r > b + RED_DOMINANCE:
            return True
    return False


def parse_exam(docx_path: Path) -> list[Part]:
    return parse_exam_document(Document(str(docx_path)))


def _line_spans(text: str) -> list[tuple[int, int]]:
    """Tra ve (start, end) cua tung dong trong text (tach boi '\\n' - xuong
    dong MEM kieu Shift+Enter cung the hien thanh '\\n' trong paragraph.text),
    da bo khoang trang dau/cuoi dong, giu nguyen vi tri thuc trong text GOC
    (de _is_red_range van tra cuu dung run mau)."""
    spans = []
    cursor = 0
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped:
            line_offset = line.index(stripped)
            start = cursor + line_offset
            spans.append((start, start + len(stripped)))
        cursor += len(line) + 1
    return spans


def _iter_body_items(document: Document):
    """Duyet cac phan tu truc tiep trong body theo DUNG THU TU xuat hien trong
    file (paragraph va bang xen ke nhau) - can thiet de biet dung "Phan/Cau"
    nao dang mo khi gap 1 bang (vd tieu de "I. LISTENING" nam TRUOC bang chua
    cau 1-4, chu khong phai trong bang)."""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def parse_exam_document(document: Document) -> list[Part]:
    # Luon co san 1 "part mac dinh" (title rong) tu dau - vi mot so de khong
    # dung dang tieu de "I. LISTENING" ma chi ghi tron "LISTENING" (khong so
    # La Ma) hoac khong co tieu de Phan nao ca. Neu bat cau hoi phai cho toi
    # khi khop duoc PART_HEADER_RE moi bat dau ghi nhan, nhung file kieu do
    # se lam MAT TRANG toan bo cau hoi (current_part luon None -> flush bi
    # bo qua) ma khong bao loi gi ca.
    default_part = Part(title="")
    parts: list[Part] = [default_part]
    current_part: Part | None = default_part
    current_question: Question | None = None
    pending_label: list[str | None] = [None]

    def flush_question() -> None:
        if current_question is not None and current_part is not None:
            current_part.questions.append(current_question)

    def process_options(full_text: str, start: int, end: int, paragraph: Paragraph) -> None:
        nonlocal current_question
        text = full_text[start:end]
        if current_question is None or not text:
            return

        # Thu nhan dien "nhieu dap an chung 1 dong" TRUOC - vi OPTION_RE (mot
        # dap an rieng dong, ".+$" greedy) se nuot ca dong nhieu dap an neu
        # thu no truoc, khien 4 dap an bi hop thanh 1.
        inline_matches = list(INLINE_OPTION_RE.finditer(full_text, start, end))
        if len(inline_matches) >= 2:
            for m in inline_matches:
                rest = m.group("rest")
                # strip() co the lam lech vi tri that cua noi dung so voi
                # m.start("rest")/end("rest") (neu rest co khoang trang dau/cuoi)
                # - tinh lai offset cho dung phan da strip.
                lstripped = rest.lstrip()
                lead = len(rest) - len(lstripped)
                content = lstripped.rstrip()
                content_start = m.start("rest") + lead
                current_question.options.append(
                    Option(
                        label=m.group("label"),
                        text=content,
                        paragraph=paragraph,
                        is_correct=_is_red_range(paragraph, m.start(), m.end()),
                        start=content_start,
                        end=content_start + len(content),
                    )
                )
            return

        # Dung substring rieng (khong dung full_text+pos) vi OPTION_RE co "^" -
        # "^" chi neo dung VI TRI 0 CUA CA CHUOI GOC, khong phai vi tri `start`,
        # nen voi full_text+pos no chi khop dong DAU TIEN cua doan (start=0),
        # cac dong sau (B, C, D... khi start>0) se luon that bai.
        option_match = OPTION_RE.match(text)
        if option_match:
            rest = option_match.group("rest")
            lstripped = rest.lstrip()
            lead = len(rest) - len(lstripped)
            content = lstripped.rstrip()
            content_start = start + option_match.start("rest") + lead
            current_question.options.append(
                Option(
                    label=option_match.group("label"),
                    text=content,
                    paragraph=paragraph,
                    is_correct=_is_red_range(paragraph, start, end),
                    start=content_start,
                    end=content_start + len(content),
                )
            )
            return

        # Dong dau tien cua paragraph (start == 0) dung numbered-list tu dong
        # cua Word: chu cai (vd "A.") do Word tu ve ra, khong co trong text
        # that - suy ra nhan theo THU TU dap an da doc duoc cua cau nay (dap
        # an dau tien -> A, tiep theo -> B,...).
        if start == 0 and current_question is not None and len(current_question.options) < len(CHOICE_LETTERS) and _has_auto_numbering(paragraph):
            current_question.options.append(
                Option(
                    label=CHOICE_LETTERS[len(current_question.options)],
                    text=text,
                    paragraph=paragraph,
                    is_correct=_is_red_range(paragraph, start, end),
                    start=start,
                    end=end,
                )
            )
            return

        current_question.paragraphs.append(paragraph)

    def process_paragraph(paragraph: Paragraph) -> None:
        nonlocal current_part, current_question
        full_text = paragraph.text
        for start, end in _line_spans(full_text):
            text = full_text[start:end]

            part_match = PART_HEADER_RE.match(text)
            if part_match:
                flush_question()
                current_question = None
                current_part = Part(title=text)
                parts.append(current_part)
                continue

            question_match = QUESTION_RE.match(text)
            if question_match:
                flush_question()
                current_question = Question(number=int(question_match.group("num")))
                # Truong hop dap an nam CHUNG dong voi "Question N." (vd
                # "Question 9.  A. Saturday   B. racket...") - xu ly phan con
                # lai cua dong ngay sau khi tao xong cau hoi, thay vi bo qua.
                process_options(full_text, start + question_match.end(), end, paragraph)
                continue

            process_options(full_text, start, end, paragraph)

    def process_table(table: Table) -> None:
        pending_label[0] = None
        for row in table.rows:
            pending_label[0] = None
            # Cell bi gop ngang bi lap lai nhieu lan trong row.cells (python-docx
            # tra ve 1 _Cell wrapper MOI cho moi cot no trai dai qua, du cung
            # tro toi 1 the <w:tc> XML - phai dedup bang _tc, khong dung == /in
            # truc tiep tren _Cell vi no khong coi 2 wrapper nay la bang nhau).
            seen_tc_ids: set[int] = set()
            unique_cells: list[_Cell] = []
            for cell in row.cells:
                if id(cell._tc) not in seen_tc_ids:
                    seen_tc_ids.add(id(cell._tc))
                    unique_cells.append(cell)

            for cell in unique_cells:
                cell_text = cell.text.strip()

                if BARE_LABEL_RE.match(cell_text):
                    pending_label[0] = cell_text
                    continue

                if pending_label[0] is not None and current_question is not None and cell_text and cell.paragraphs:
                    content_paragraph = cell.paragraphs[0]
                    raw = content_paragraph.text
                    content_start = raw.find(cell_text)
                    if content_start == -1:
                        content_start = 0
                    current_question.options.append(
                        Option(
                            label=pending_label[0],
                            text=cell_text,
                            paragraph=content_paragraph,
                            is_correct=_is_red_range(content_paragraph, 0, len(raw)),
                            start=content_start,
                            end=content_start + len(cell_text),
                        )
                    )
                    pending_label[0] = None
                    continue

                for cell_paragraph in cell.paragraphs:
                    process_paragraph(cell_paragraph)

    for item in _iter_body_items(document):
        if isinstance(item, Table):
            process_table(item)
        else:
            process_paragraph(item)

    flush_question()
    return parts


def validate_exam(parts: list[Part]) -> list[str]:
    issues: list[str] = []
    all_numbers = sorted(q.number for part in parts for q in part.questions)

    if not all_numbers:
        issues.append(
            "Khong doc duoc cau hoi nao tu file. Kiem tra dinh dang: cau hoi phai "
            "ghi dang \"Question 1.\" hoac \"Cau 1.\" (co dau cham/hai cham sau so)."
        )
        return issues

    # Luon gia dinh de bat dau tu cau 1 - khong dung all_numbers[0] lam moc
    # dau, vi neu ca 1 khoang cau dau (vi du cau 1-4) mat hoan toan (nam
    # trong bang, khong doc duoc), all_numbers[0] se la so cua cau DAU TIEN
    # DOC DUOC (vi du 5), khien khoang thieu truoc do bi lot mat.
    expected = set(range(1, all_numbers[-1] + 1))
    missing = sorted(expected - set(all_numbers))
    if missing:
        issues.append(
            f"Khong tim thay cau: {', '.join(str(n) for n in missing)} "
            "(co the do nam trong bang - xem canh bao 'file co N bang' va sua file goc, bo bang)"
        )

    for part in parts:
        for question in part.questions:
            if not question.options:
                issues.append(f"Cau {question.number}: khong doc duoc dap an nao")
            elif not any(opt.is_correct for opt in question.options):
                issues.append(f"Cau {question.number}: khong xac dinh duoc dap an dung (khong thay dap an nao to mau do)")

    return issues
